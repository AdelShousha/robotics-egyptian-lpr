#!/usr/bin/env python3
"""
Convert markdown documentation files to a Microsoft Word document.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Files to compile in order
MD_FILES = [
    '01-overview.md',
    '02-dataset-preparation.md',
    '03-model-training.md',
    '04-backend-api.md',
    '05-dashboard.md',
    '06-hardware-setup.md',
    '07-deployment.md',
    '08-system-architecture.md'
]

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_table(lines, start_idx):
    """Parse a markdown table and return table data and end index."""
    table_lines = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        table_lines.append(lines[i])
        i += 1

    if len(table_lines) < 2:
        return None, start_idx

    # Parse header
    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

    # Skip separator line
    data_start = 2 if len(table_lines) > 1 and '---' in table_lines[1] else 1

    # Parse data rows
    rows = []
    for line in table_lines[data_start:]:
        if '|' in line:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if row:
                rows.append(row)

    return {'header': header, 'rows': rows}, i - 1

def add_table_to_doc(doc, table_data):
    """Add a formatted table to the document."""
    if not table_data or not table_data['header']:
        return

    num_cols = len(table_data['header'])
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add header
    header_row = table.rows[0]
    for i, cell_text in enumerate(table_data['header']):
        cell = header_row.cells[i]
        cell.text = cell_text
        set_cell_shading(cell, 'E0E0E0')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Add data rows
    for row_data in table_data['rows']:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i < len(row.cells):
                cell = row.cells[i]
                # Handle bold text
                if cell_text.startswith('**') and cell_text.endswith('**'):
                    cell_text = cell_text[2:-2]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(cell_text)
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    doc.add_paragraph()  # Space after table

def add_code_block(doc, code_text):
    """Add a formatted code block to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_image_placeholder(doc, alt_text, caption):
    """Add an image placeholder box to the document."""
    # Add placeholder box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)

    run = p.add_run(f"[IMAGE PLACEHOLDER]\n{alt_text}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    # Add caption if present
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.italic = True
        cap_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

def process_inline_formatting(paragraph, text):
    """Process inline formatting (bold, italic, code, links)."""
    # Pattern for different formatting
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),
        (r'\*(.+?)\*', 'italic'),
        (r'`([^`]+)`', 'code'),
        (r'\[([^\]]+)\]\(([^)]+)\)', 'link'),
    ]

    # Simple approach: just add the text with basic formatting
    # Remove markdown link syntax
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    # Process bold
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            # Process italic within non-bold parts
            sub_parts = re.split(r'(\*[^*]+\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*') and not sub_part.startswith('**'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                elif sub_part:
                    # Process inline code
                    code_parts = re.split(r'(`[^`]+`)', sub_part)
                    for code_part in code_parts:
                        if code_part.startswith('`') and code_part.endswith('`'):
                            run = paragraph.add_run(code_part[1:-1])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                        elif code_part:
                            paragraph.add_run(code_part)

def convert_md_to_docx(md_files, output_path):
    """Convert markdown files to a Word document."""
    doc = Document()

    # Set document title
    title = doc.add_heading('Egyptian License Plate Recognition System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].italic = True

    doc.add_paragraph()
    add_horizontal_line(doc)

    # Add table of contents placeholder
    toc_heading = doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. Dataset Preparation',
        '3. Model Training',
        '4. Backend API',
        '5. Dashboard',
        '6. Hardware Setup',
        '7. Deployment',
        '8. System Architecture'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_page_break()

    # Process each markdown file
    for md_file in md_files:
        file_path = os.path.join(os.path.dirname(__file__), md_file)

        if not os.path.exists(file_path):
            print(f"Warning: {md_file} not found, skipping...")
            continue

        print(f"Processing {md_file}...")

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_block_content = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    add_code_block(doc, '\n'.join(code_block_content))
                    code_block_content = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # Handle headers
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 4)

            # Handle horizontal rules
            elif line.strip() == '---':
                add_horizontal_line(doc)

            # Handle images
            elif line.startswith('!['):
                match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if match:
                    alt_text = match.group(1)
                    # Check for caption on next line
                    caption = ''
                    if i + 1 < len(lines) and lines[i + 1].startswith('*Caption:'):
                        caption = lines[i + 1].strip('*')
                        i += 1
                    add_image_placeholder(doc, alt_text, caption)

            # Handle tables
            elif '|' in line and not line.startswith('```'):
                table_data, end_idx = parse_table(lines, i)
                if table_data:
                    add_table_to_doc(doc, table_data)
                    i = end_idx

            # Handle bullet points
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                process_inline_formatting(p, text)

            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(style='List Number')
                process_inline_formatting(p, text)

            # Handle regular paragraphs
            elif line.strip() and not line.startswith('*Caption:'):
                p = doc.add_paragraph()
                process_inline_formatting(p, line)

            i += 1

        # Add page break between sections (except last)
        if md_file != md_files[-1]:
            doc.add_page_break()

    # Save document
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_file = os.path.join(script_dir, 'Egyptian_LPR_Documentation.docx')
    convert_md_to_docx(MD_FILES, output_file)
