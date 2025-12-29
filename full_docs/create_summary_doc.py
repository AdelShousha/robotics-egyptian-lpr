#!/usr/bin/env python3
"""
Convert condensed markdown documentation to Microsoft Word document.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_table(lines, start_idx):
    """Parse a markdown table and return table data and end index."""
    table_lines = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        table_lines.append(lines[i])
        i += 1

    if len(table_lines) < 2:
        return None, start_idx

    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    data_start = 2 if len(table_lines) > 1 and '---' in table_lines[1] else 1

    rows = []
    for line in table_lines[data_start:]:
        if '|' in line:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if row:
                rows.append(row)

    return {'header': header, 'rows': rows}, i - 1

def add_table_to_doc(doc, table_data):
    """Add a formatted table to the document."""
    if not table_data or not table_data['header']:
        return

    num_cols = len(table_data['header'])
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = table.rows[0]
    for i, cell_text in enumerate(table_data['header']):
        cell = header_row.cells[i]
        cell.text = cell_text
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_data in table_data['rows']:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i < len(row.cells):
                cell = row.cells[i]
                if cell_text.startswith('**') and cell_text.endswith('**'):
                    cell_text = cell_text[2:-2]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(cell_text)
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    doc.add_paragraph()

def add_code_block(doc, code_text):
    """Add a formatted code block to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def add_image_placeholder(doc, alt_text, caption):
    """Add an image placeholder box to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)

    run = p.add_run(f"[IMAGE: {alt_text}]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.italic = True
        cap_run.font.color.rgb = RGBColor(100, 100, 100)

def process_inline_formatting(paragraph, text):
    """Process inline formatting."""
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            code_parts = re.split(r'(`[^`]+`)', part)
            for code_part in code_parts:
                if code_part.startswith('`') and code_part.endswith('`'):
                    run = paragraph.add_run(code_part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif code_part:
                    paragraph.add_run(code_part)

def convert_md_to_docx(md_file, output_path):
    """Convert markdown file to a Word document."""
    doc = Document()

    # Title
    title = doc.add_heading('Egyptian License Plate Recognition System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Project Documentation Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].italic = True

    add_horizontal_line(doc)

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_block_content))
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            run.bold = True
            run.font.size = Pt(11)

        # Horizontal rules
        elif line.strip() == '---':
            add_horizontal_line(doc)

        # Images
        elif line.startswith('!['):
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                caption = ''
                if i + 1 < len(lines) and lines[i + 1].startswith('*Caption:'):
                    caption = lines[i + 1].strip('*')
                    i += 1
                add_image_placeholder(doc, alt_text, caption)

        # Tables
        elif '|' in line and not line.startswith('```'):
            table_data, end_idx = parse_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
                i = end_idx

        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)

        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)

        # Regular paragraphs
        elif line.strip() and not line.startswith('*Caption:'):
            p = doc.add_paragraph()
            process_inline_formatting(p, line)

        i += 1

    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, 'Egyptian_LPR_Summary.md')
    output_file = os.path.join(script_dir, 'Egyptian_LPR_Summary.docx')
    convert_md_to_docx(md_file, output_file)
